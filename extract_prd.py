import docx
import os

doc_path = r"d:\placement\placement analysis doc\Placement_Internship_Analytics_PRD.docx"
if os.path.exists(doc_path):
    doc = docx.Document(doc_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            full_text.append(" | ".join(row_text))
            
    print("\n".join(full_text))
else:
    print(f"File not found: {doc_path}")
